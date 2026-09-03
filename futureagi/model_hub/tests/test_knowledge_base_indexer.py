"""Unit coverage for ``KBIndexer`` — the knowledge-base ingestion engine.

``model_hub/utils/kb_indexer.py`` is 456 lines of document parsing, text
cleaning, chunking, embedding dispatch and vector search, live in three call
sites in ``model_hub/tasks/develop_dataset.py`` plus the ee agenthub. It had no
tests: ``test_kb_helpers.py`` covers the scheduling/cancellation orchestration
*around* it, never the indexer itself.

Both external boundaries are mocked, so nothing here reaches the embedding
service or object storage:
  * ``EmbeddingManager``  — embedding + vector-store calls
  * ``get_storage_client`` — MinIO/S3
"""

import os
import tempfile

import pytest

from model_hub.utils.kb_indexer import (
    KB_INDEX_COL_NAME,
    KB_TABLE_NAME,
    Chunk,
    KBIndexer,
)

MODULE = "model_hub.utils.kb_indexer"


@pytest.fixture
def indexer(mocker):
    """A ``KBIndexer`` with both external boundaries replaced by mocks."""
    embedding_manager = mocker.MagicMock()
    mocker.patch(f"{MODULE}.EmbeddingManager", return_value=embedding_manager)
    storage = mocker.MagicMock()
    mocker.patch(f"{MODULE}.get_storage_client", return_value=storage)

    kb_indexer = KBIndexer()
    kb_indexer._test_embedding_manager = embedding_manager
    kb_indexer._test_storage = storage
    return kb_indexer


def _write(tmp_path, name, content, mode="w", **kwargs):
    path = tmp_path / name
    if mode == "w":
        path.write_text(content, **kwargs)
    else:
        path.write_bytes(content)
    return str(path)


class TestCleanText:
    """``_clean_text`` is pure — no mocking needed."""

    def test_empty_input_returns_empty_string(self, indexer):
        assert indexer._clean_text("") == ""
        assert indexer._clean_text(None) == ""

    def test_null_bytes_are_stripped(self, indexer):
        assert "\x00" not in indexer._clean_text("he\x00llo")
        assert indexer._clean_text("he\x00llo") == "hello"

    def test_whitespace_is_normalised_and_trimmed(self, indexer):
        assert indexer._clean_text("  a \n\n b\t\tc  ") == "a b c"

    def test_cid_artifacts_are_removed(self, indexer):
        assert indexer._clean_text("start(cid:123)end") == "startend"

    def test_ordinary_text_is_preserved(self, indexer):
        assert indexer._clean_text("Chapter 1. The beginning.") == (
            "Chapter 1. The beginning."
        )


class TestFileReaders:
    """The per-extension readers, exercised against real temp files."""

    def test_process_text_returns_file_contents(self, indexer, tmp_path):
        path = _write(tmp_path, "doc.txt", "hello knowledge base")

        assert indexer.process_text(path) == "hello knowledge base"

    def test_process_text_returns_none_for_empty_file(self, indexer, tmp_path):
        path = _write(tmp_path, "empty.txt", "")

        assert indexer.process_text(path) is None

    def test_process_docx_joins_paragraphs(self, indexer, tmp_path):
        import docx as docx_module

        document = docx_module.Document()
        document.add_paragraph("first para")
        document.add_paragraph("second para")
        path = str(tmp_path / "doc.docx")
        document.save(path)

        text = indexer.process_docx(path)

        assert "first para" in text
        assert "second para" in text

    def test_process_rtf_extracts_plain_text(self, indexer, tmp_path):
        path = _write(
            tmp_path, "doc.rtf", r"{\rtf1\ansi hello from rtf\par}"
        )

        assert "hello from rtf" in indexer.process_rtf(path)

    def test_process_pdf_cleans_extracted_pages(self, indexer, mocker):
        page_one = mocker.MagicMock(page_content="page  one\x00")
        page_two = mocker.MagicMock(page_content="page two(cid:7)")
        mocker.patch(
            f"{MODULE}.PyPDFLoader",
            return_value=mocker.MagicMock(load=lambda: [page_one, page_two]),
        )

        text = indexer.process_pdf("/tmp/whatever.pdf")

        assert text == "page one page two"

    def test_process_pdf_propagates_loader_failure(self, indexer, mocker):
        mocker.patch(f"{MODULE}.PyPDFLoader", side_effect=RuntimeError("corrupt"))

        with pytest.raises(RuntimeError):
            indexer.process_pdf("/tmp/broken.pdf")


class TestProcessFileDispatch:
    """``process_file`` routes on the file extension."""

    @pytest.mark.parametrize(
        "filename,reader",
        [
            ("doc.pdf", "process_pdf"),
            ("doc.txt", "process_text"),
            ("doc.docx", "process_docx"),
            ("doc.rtf", "process_rtf"),
        ],
    )
    def test_extension_selects_the_matching_reader(
        self, mocker, indexer, filename, reader
    ):
        mocker.patch.object(indexer, reader, return_value="extracted text")
        process_content = mocker.patch.object(indexer, "process_content")

        indexer.process_file(f"/tmp/{filename}", "file-1", "kb-1", "org-1")

        process_content.assert_called_once_with(
            "extracted text", "file-1", "kb-1", "org-1"
        )

    def test_empty_extraction_raises_value_error(self, mocker, indexer):
        mocker.patch.object(indexer, "process_text", return_value=None)
        process_content = mocker.patch.object(indexer, "process_content")

        with pytest.raises(ValueError, match="No content extracted"):
            indexer.process_file("/tmp/doc.txt", "file-1", "kb-1", "org-1")

        process_content.assert_not_called()

    @pytest.mark.parametrize("filename", ["doc.csv", "doc.PDF", "doc"])
    def test_unsupported_extension_fails_rather_than_silently_passing(
        self, mocker, indexer, filename
    ):
        """An unroutable extension fails loudly rather than indexing nothing.

        No ``else`` branch assigns ``text``, so the guard raises
        UnboundLocalError. The dispatch is case-sensitive, so ``.PDF`` takes the
        same path as ``.csv``.
        """
        process_content = mocker.patch.object(indexer, "process_content")

        with pytest.raises((UnboundLocalError, ValueError)):
            indexer.process_file(f"/tmp/{filename}", "file-1", "kb-1", "org-1")

        process_content.assert_not_called()


class TestProcessContent:
    """Chunking and the batched embedding dispatch."""

    def test_chunks_are_built_with_stable_ids_and_metadata(self, indexer):
        indexer.process_content("word " * 500, "file-1", "kb-1", "org-1")

        assert indexer.chunks
        first = indexer.chunks[0]
        assert isinstance(first, Chunk)
        assert first.file_id == "file-1"
        assert first.organization_id == "org-1"
        assert first.chunk_id.startswith("file-1_")
        assert first.metadata["chunk_type"] == "semantic"

    def test_embedding_manager_receives_kb_id_and_table(self, indexer):
        indexer.process_content("word " * 500, "file-1", "kb-1", "org-1")

        call = indexer._test_embedding_manager.parallel_process_metadata.call_args
        assert call.kwargs["eval_id"] == "kb-1"
        assert call.kwargs["table_name"] == KB_TABLE_NAME
        assert call.kwargs["inputs_formater"] == [KB_INDEX_COL_NAME]
        assert all(
            KB_INDEX_COL_NAME in metadata for metadata in call.kwargs["metadatas"]
        )

    def test_short_text_produces_a_single_batch(self, indexer):
        indexer.process_content("a short document", "file-1", "kb-1", "org-1")

        assert (
            indexer._test_embedding_manager.parallel_process_metadata.call_count == 1
        )
        assert len(indexer.chunks) == 1

    def test_batch_failure_is_raised_not_swallowed(self, indexer):
        indexer._test_embedding_manager.parallel_process_metadata.side_effect = (
            RuntimeError("vector store down")
        )

        with pytest.raises(RuntimeError, match="Errors occurred during batch"):
            indexer.process_content("word " * 500, "file-1", "kb-1", "org-1")

        assert indexer.chunks == []

    def test_chunks_accumulate_across_calls(self, indexer):
        indexer.process_content("first document", "file-1", "kb-1", "org-1")
        indexer.process_content("second document", "file-2", "kb-1", "org-1")

        file_ids = {chunk.file_id for chunk in indexer.chunks}
        assert file_ids == {"file-1", "file-2"}


class TestS3Download:
    def test_object_key_is_scoped_to_the_knowledge_base(self, indexer, tmp_path):
        local = indexer.download_s3_file(
            "prefix", "org-1", "kb-1", "file-1", str(tmp_path), "pdf"
        )

        indexer._test_storage.fget_object.assert_called_once()
        bucket, object_key, local_path = indexer._test_storage.fget_object.call_args[0]
        assert object_key == "knowledge-base/kb-1/file-1.pdf"
        assert local_path == os.path.join(str(tmp_path), "file-1.pdf")
        assert local == local_path

    def test_download_failure_propagates(self, indexer, tmp_path):
        indexer._test_storage.fget_object.side_effect = OSError("no such key")

        with pytest.raises(OSError):
            indexer.download_s3_file(
                "prefix", "org-1", "kb-1", "file-1", str(tmp_path), "pdf"
            )


class TestProcessS3File:
    def test_success_returns_file_and_kb_ids(self, mocker, indexer):
        mocker.patch.object(indexer, "download_s3_file", return_value="/tmp/f.txt")
        mocker.patch.object(indexer, "process_file")

        result = indexer.process_s3_file("docs/f.txt", "file-1", "kb-1", "org-1")

        assert result == {"file_id": "file-1", "kb_id": "kb-1"}

    def test_extension_is_taken_from_the_source_path(self, mocker, indexer):
        download = mocker.patch.object(
            indexer, "download_s3_file", return_value="/tmp/f.docx"
        )
        mocker.patch.object(indexer, "process_file")

        indexer.process_s3_file("docs/report.docx", "file-1", "kb-1", "org-1")

        assert download.call_args[0][-1] == "docx"

    def test_failure_is_reported_as_an_error_dict_not_raised(self, mocker, indexer):
        mocker.patch.object(
            indexer, "download_s3_file", side_effect=OSError("gone")
        )

        result = indexer.process_s3_file("docs/f.txt", "file-1", "kb-1", "org-1")

        assert result["file_id"] == "file-1"
        assert result["kb_id"] == "kb-1"
        assert "gone" in result["error"]

    def test_local_source_file_is_removed_afterwards(self, mocker, indexer):
        handle, path = tempfile.mkstemp(suffix=".txt")
        os.close(handle)
        mocker.patch.object(indexer, "download_s3_file", return_value=path)
        mocker.patch.object(indexer, "process_file")

        indexer.process_s3_file(path, "file-1", "kb-1", "org-1")

        assert not os.path.isfile(path)


class TestSearchAndDelete:
    """Thin delegations — assert the arguments handed to the embedding manager."""

    def test_search_queries_the_kb_table(self, indexer):
        indexer._test_embedding_manager.retrieve_avg_rag_based_examples.return_value = [
            {"chunk_text": "hit"}
        ]

        results = indexer.search("kb-1", "what is x")

        assert results == [{"chunk_text": "hit"}]
        kwargs = (
            indexer._test_embedding_manager
            .retrieve_avg_rag_based_examples.call_args.kwargs
        )
        assert kwargs["eval_id"] == "kb-1"
        assert kwargs["table_name"] == KB_TABLE_NAME
        assert kwargs["input_cols"] == [KB_INDEX_COL_NAME]

    def test_get_subset_kb_id_delegates_with_index_col_type(self, indexer):
        indexer._test_embedding_manager.get_relevant_chunks.return_value = "kb-subset"

        assert indexer.get_subset_kb_id("query", "kb-1") == "kb-subset"
        kwargs = indexer._test_embedding_manager.get_relevant_chunks.call_args.kwargs
        assert kwargs["eval_id"] == "kb-1"
        assert kwargs["table_name"] == KB_TABLE_NAME

    def test_get_data_subset_passes_top_k_and_threshold(self, indexer):
        indexer.get_data_subset_kb_id(["a", "b"], "kb-1", top_k=7)

        kwargs = (
            indexer._test_embedding_manager
            .retrieve_avg_rag_based_examples.call_args.kwargs
        )
        assert kwargs["top_k"] == 7
        assert kwargs["threshold"] == 0.35
        # one input_col per query
        assert kwargs["input_cols"] == [KB_INDEX_COL_NAME] * 2

    def test_remove_chunks_returns_ids_on_success(self, indexer):
        result = indexer.remove_chunks_from_kb("file-1", "kb-1", "org-1")

        assert result == {"file_id": "file-1", "kb_id": "kb-1"}
        indexer._test_embedding_manager.delete_chunks.assert_called_once_with(
            "file-1", "kb-1", KB_TABLE_NAME, "org-1"
        )

    def test_remove_chunks_reports_failure_as_error_dict(self, indexer):
        indexer._test_embedding_manager.delete_chunks.side_effect = RuntimeError(
            "delete failed"
        )

        result = indexer.remove_chunks_from_kb("file-1", "kb-1", "org-1")

        assert "delete failed" in result["error"]
